
import difflib
import sys
import json
import unicodedata
import ast
import time
import uuid
from openai import OpenAI
from docx import Document
from dotenv import load_dotenv
from docx.shared import Pt

# Ngarko variablat nga .env file (opsionale)
load_dotenv()

# Merr API key
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
  raise ValueError("OPENAI_API_KEY nuk është vendosur. Vendosni çelësin në .env ose në variablat e mjedisit.")
client = OpenAI(api_key=api_key)


def normalize_key(s: str) -> str:
  """Normalize key: remove diacritics, lowercase, replace non-alnum with underscore."""
  s = s or ""
  s = unicodedata.normalize('NFKD', s)
  s = s.encode('ascii', 'ignore').decode()
  s = s.lower()
  s = re.sub(r"[^a-z0-9]+", "_", s)
  return s.strip('_')


def extract_json(text: str) -> dict:
  """Try to extract a JSON object from text and parse it."""
  text = text.strip()
  try:
    return json.loads(text)
  except Exception:
    m = re.search(r"\{[\s\S]*\}", text)
    if not m:
      raise ValueError(f"Përgjigjja e AI nuk përmban JSON të vlefshëm. Përgjigjja: {text}")
    try:
      return json.loads(m.group(0))
    except Exception as e:
      raise ValueError(f"Nuk arrita të deshifroj JSON nga përgjigjja e AI: {e}\nPërgjigjja: {text}")


def fill_template_from_data(data: dict, template_path: str = 'shabllon.docx', out_path: str = 'Ditari_Final.docx'):
  """Replace placeholders {{key}} in Word template using normalized keys from data."""
  # helper to format/normalize values for specific keys
  def format_value(key_norm: str, val, tema_text: str):
    if val is None:
      return ''
    # convert list/tuple values to comma-separated string
    if isinstance(val, (list, tuple)):
      return ', '.join([str(x) for x in val])
    # if string looks like a list literal, try to parse it
    if isinstance(val, str):
      s = val.strip()
      if s.startswith('[') and s.endswith(']'):
        try:
          parsed = json.loads(s)
          if isinstance(parsed, (list, tuple)):
            return ', '.join([str(x) for x in parsed])
        except Exception:
          try:
            parsed = ast.literal_eval(s)
            if isinstance(parsed, (list, tuple)):
              return ', '.join([str(x) for x in parsed])
          except Exception:
            pass

    # default string
    out = str(val)

    # For keyword fields, ensure nicer formatting and add extra guidance
    if 'fjalet' in key_norm:
      # remove surrounding brackets if present and ensure comma-separated
      return re.sub(r'[\[\]]', '', out).strip()

    # Enrich certain pedagogical fields with more detail and drawing hints
    enrich_keys = ['lidhja_e_temes_me_njohurite_e_meparshme', 'ndertimi_i_njohurive', 'perforcimi_i_te_nxenit']
    if key_norm in enrich_keys:
      extra = ''
      # if short, add suggested activities and assessment prompts
      if len(out) < 80:
        extra += ' Shtoni: (1) Pyetje hyrëse për aktivizim, (2) Aktivitet praktik me shembuj, (3) Diskutim/ndarje grupesh.'
      else:
        extra += ' Përmbledhje: shfrytëzoni aktivitete praktike, pyetje kontrolluese dhe ushtrime aplikuese.'

      # suggest simple drawing if theme or text indicates geometry/figures
      geom_keywords = ['gjeometr', 'trekënd', 'triang', 'trigonometri', 'figura', 'forma', 'shkronj']
      combined = (tema_text or '') + ' ' + out
      if any(k in combined.lower() for k in geom_keywords):
        extra += ' Vizatoni në tabelë: trekëndësh të drejtë (katetet a, b; hipotenuza c) dhe etiketoji.'
        extra += '\n[Shembull vizatimi ASCII]\n   |\n   | \\\n+   |  \\\n+   |___\\\\\n'

      return out + extra

    return out

  tema_text = data.get('tema', '')
  norm_data = {normalize_key(k): format_value(normalize_key(k), v, tema_text) for k, v in data.items()}

  # Normalize template path
  base_dir = os.path.dirname(os.path.abspath(__file__))
  if template_path.endswith('.docx.docx'):
    template_path = template_path.replace('.docx.docx', '.docx')

  # Try absolute path in project folder
  template_abs = os.path.join(base_dir, template_path)
  if os.path.exists(template_abs):
    template_path = template_abs
  else:
    # Fallback to default template name in project folder
    fallback = os.path.join(base_dir, 'shabllon.docx')
    if os.path.exists(fallback):
      template_path = fallback

  # Open template
  doc = Document(template_path)
  placeholder_re = re.compile(r"{{\s*(.*?)\s*}}")
  def find_value_for_key(key_norm: str):
    if key_norm in norm_data:
      return norm_data[key_norm]
    for k in norm_data:
      if key_norm.startswith(k) or k.startswith(key_norm) or key_norm in k or k in key_norm:
        return norm_data[k]
    matches = difflib.get_close_matches(key_norm, list(norm_data.keys()), n=1, cutoff=0.6)
    if matches:
      return norm_data[matches[0]]
    return None

  for table in doc.tables:
    for row in table.rows:
      for cell in row.cells:
        # Process each paragraph and preserve original runs/formatting.
        for para in cell.paragraphs:
          # Quick skip if no placeholder
          if not placeholder_re.search(para.text):
            continue
          # Replace placeholders at run-level to preserve original fonts.
          full = ''.join([r.text for r in para.runs])
          matches = list(placeholder_re.finditer(full))
          if not matches:
            continue

          # helper to replace span [s,e) in runs with replacement text
          def replace_span_in_runs(paragraph, s, e, replacement):
            runs = paragraph.runs
            # compute run index and offset for s and e
            acc = 0
            start_run = start_off = end_run = end_off = None
            for i, r in enumerate(runs):
              ln = len(r.text)
              if start_run is None and acc + ln > s:
                start_run = i
                start_off = s - acc
              if end_run is None and acc + ln >= e:
                end_run = i
                end_off = e - acc
                break
              acc += ln
            if start_run is None or end_run is None:
              return
            # build new text for start run: prefix + replacement + suffix
            prefix = runs[start_run].text[:start_off]
            suffix = runs[end_run].text[end_off:]
            runs[start_run].text = prefix + replacement + suffix
            # clear texts of intermediate runs
            for j in range(start_run + 1, end_run + 1):
              runs[j].text = ''

          # process matches in reverse so indices stay valid
          for m in reversed(matches):
            key_raw = m.group(1)
            key_norm = normalize_key(key_raw)
            value = find_value_for_key(key_norm)
            if value is None:
              continue
            replace_span_in_runs(para, m.start(), m.end(), str(value))

  try:
    doc.save(out_path)
    return out_path
  except PermissionError:
    # try an alternate name, then fall back to a unique timestamped filename
    alt = out_path.replace('.docx', '_new.docx')
    try:
      doc.save(alt)
      return alt
    except PermissionError:
      uniq = out_path.replace('.docx', f'_{int(time.time())}_{uuid.uuid4().hex[:6]}.docx')
      doc.save(uniq)
      return uniq


def request_plan_json(tema: str) -> dict:
  """Ask the model to return a single JSON object with predefined keys for the lesson plan."""
  prompt = f"""
Gjenero një plan ditor mësimi për temën: {tema}.
Përgjigju VETËM me një objekt JSON (pa tekst tjetër) me këto çelësa:
fusha, lenda, shkalla, klasa, tema, tema_2, situata, lidhja, burimet, fjalet_kryesore, metodologjia,
lidhja_e_temes_me_njohurite_e_meparshme, ndertimi_i_njohurive, perforcimi_i_te_nxenit,
rezultatet, shenime_vleresuese, detyra_shtepie
"""

  resp = client.chat.completions.create(
    model='gpt-3.5-turbo',
    messages=[{"role": "user", "content": prompt}],
    temperature=0.2,
  )
  text = resp.choices[0].message.content
  return extract_json(text)


def krijo_ditarin(tema_kerkuar: str):
  print(f"AI po gjeneron ditarin për: {tema_kerkuar}...")
  data = request_plan_json(tema_kerkuar)
  # Ensure tema field exists
  if 'tema' not in {normalize_key(k): v for k, v in data.items()}:
    data['tema'] = tema_kerkuar
  # Ensure 'arsimi' mirrors 'lenda' when appropriate
  if 'arsimi' not in data and 'lenda' in data:
    data['arsimi'] = data.get('lenda')

  # Build structured 'shenime_vleresuese' with 3 evaluation notes (N2,N3,N4)
  rezultate_text = data.get('rezultatet') or data.get('rezultatet'.replace('_',' ')) or ''
  tema_text = data.get('tema') or tema_kerkuar
  sh1 = f"N2: Nxënësi kryen veprime për të arritur rezultatin e pritshëm ({rezultate_text})." if rezultate_text else "N2: Nxënësi kryen veprime të përshtatshme për të demonstruar zotësitë e pritura."
  sh2 = f"N3: Nxënësi njohon dhe përshkruan masat/konceptet e nevojshme lidhur me temën '{tema_text}'."
  sh3 = f"N4: Nxënësi përdor vetinë për të zgjidhur detyra të ngjashme dhe tregon aplikim praktik." 
  data['shenime_vleresuese'] = "\n".join([sh1, sh2, sh3])

  out = fill_template_from_data(data)
  print(f"Sukses! Ruajtur: {out}")


if __name__ == '__main__':
  if len(sys.argv) > 1:
    tema = ' '.join(sys.argv[1:])
  else:
    tema = input('Shkruaj temën (p.sh. "Matematika 10 - Teorema e Pitagorës"): ').strip()
  krijo_ditarin(tema)